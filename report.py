# Creates: PBL-IV REPORT.docx with cover, certificate, declaration, abstract, ack,
# TOC (roman i–iv), chapters (arabic 1..), headers/footers, and margins.

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date

# ---------- CONFIGURE THESE ----------
TITLE = "AGENTIC FRAGMENTED LOGISTICS MANAGEMENT SYSTEM"
SUBTITLE = "A Natural Language Powered Logistics Orchestration Platform"
DEGREE = "Bachelor of Engineering"
BRANCH = "Computer Engineering"
STUDENT_NAME = "[Your Name]"
ROLL_NO = "[Roll Number]"
COLLEGE = "[Your College Name]"
UNIVERSITY = "[University Name]"
GUIDE_NAME = "[Guide Name]"
HOD_NAME = "[HOD Name]"
PRINCIPAL_NAME = "[Principal Name]"
ACADEMIC_YEAR = "2024–2025"
CITY = "[City]"
TODAY = date.today().strftime("%d/%m/%Y")
# -------------------------------------

def set_margins(section, all_inches=1.0):
    v = Inches(all_inches)
    section.top_margin = v
    section.bottom_margin = v
    section.left_margin = v
    section.right_margin = v

def set_base_styles(doc: Document):
    styles = doc.styles
    # Normal
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # Title
    t = styles["Title"]
    t.font.name = "Times New Roman"
    t.font.size = Pt(18)
    t.font.bold = True

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True

def _add_page_number_footer(section, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    # Footer with PAGE field
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = alignment
    run = p.add_run()

    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldSeparate = OxmlElement('w:fldChar')
    fldSeparate.set(qn('w:fldCharType'), 'separate')
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')

    run._r.append(fldBegin)
    run._r.append(instrText)
    run._r.append(fldSeparate)
    run._r.append(fldEnd)

def _set_page_number_format(section, fmt: str = "roman", start_at: int = 1):
    # fmt: "roman" | "decimal"
    sectPr = section._sectPr
    # remove existing pgNumType if any
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pgNum = OxmlElement('w:pgNumType')
    pgNum.set(qn('w:fmt'), 'roman' if fmt == 'roman' else 'decimal')
    pgNum.set(qn('w:start'), str(start_at))
    sectPr.append(pgNum)

def _insert_toc(paragraph):
    # Inserts a TOC field: \o 1-3 levels, \h hyperlinks, \z hides page numbers in web, \u
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldSeparate = OxmlElement('w:fldChar'); fldSeparate.set(qn('w:fldCharType'), 'separate')
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instrText); run._r.append(fldSeparate); run._r.append(fldEnd)

def add_cover(doc: Document):
    section = doc.sections[0]
    set_margins(section, 1.0)
    section.different_first_page_header_footer = True  # no number on cover
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE + "\n")
    r.bold = True; r.font.size = Pt(20); r.font.name = "Times New Roman"
    p.add_run(SUBTITLE + "\n").italic = True
    doc.add_paragraph().add_run("\n")

    # Project Report lines
    for line in [
        "Project Report",
        "Submitted in partial fulfillment of the requirements for the degree of",
        DEGREE,
        "in",
        BRANCH,
        "",
        "Submitted by:",
        f"{STUDENT_NAME}",
        f"Roll No.: {ROLL_NO}",
        "",
        f"{COLLEGE}",
        f"{UNIVERSITY}",
        "",
        f"Academic Year: {ACADEMIC_YEAR}",
    ]:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run(line)

    doc.add_page_break()

def add_certificate_decl_frontmatter(doc: Document):
    # New section for roman numbering front matter
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(sec, 1.0)
    _set_page_number_format(sec, "roman", 1)
    _add_page_number_footer(sec)

    # Certificate
    doc.add_paragraph("CERTIFICATE", style="Heading 1")
    t = doc.add_paragraph()
    t.add_run(
        f'This is to certify that the project report titled "{TITLE}" submitted by {STUDENT_NAME} (Roll No. {ROLL_NO}) '
        f'in partial fulfillment for the award of the {DEGREE} in {BRANCH} for the academic year {ACADEMIC_YEAR} '
        f'is a bonafide record of work carried out under my supervision.'
    )
    doc.add_paragraph("\n")
    grid = [
        ("Guide:", GUIDE_NAME),
        ("Head of Department:", HOD_NAME),
        ("Principal:", PRINCIPAL_NAME),
        ("Date:", TODAY),
    ]
    for k,v in grid:
        p = doc.add_paragraph()
        p.add_run(f"{k} ").bold = True
        p.add_run(v)
    doc.add_page_break()

    # Declaration
    doc.add_paragraph("DECLARATION", style="Heading 1")
    d = doc.add_paragraph()
    d.add_run(
        f'I, {STUDENT_NAME} (Roll No. {ROLL_NO}), hereby declare that the project work entitled "{TITLE}" '
        f'submitted to {UNIVERSITY} is my original work carried out under the supervision of {GUIDE_NAME}. '
        f'This work has not been submitted elsewhere for any degree or diploma.'
    )
    doc.add_paragraph(f"\nPlace: {CITY}")
    doc.add_paragraph(f"Date: {TODAY}")
    doc.add_paragraph(f"\nSignature: ______________________")
    doc.add_paragraph(f"Name: {STUDENT_NAME}")
    doc.add_page_break()

    # Abstract
    doc.add_paragraph("ABSTRACT", style="Heading 1")
    doc.add_paragraph(
        "This project implements an agentic logistics management workflow using a LangGraph pipeline with an LLM-based "
        "planner for intent classification and deterministic tools for execution. It supports conversational CRUD for "
        "drivers, vehicles, trips, loads, expenses, and locations via web chat and WhatsApp (Twilio), with multi-turn "
        "context (focus) and guided completion for missing fields."
    )
    doc.add_page_break()

    # Acknowledgement
    doc.add_paragraph("ACKNOWLEDGEMENT", style="Heading 1")
    doc.add_paragraph(
        "I thank my guide, department, and institution for guidance and facilities. I also acknowledge open-source "
        "communities behind LangGraph, Twilio, SQLAlchemy, FastAPI, and Python."
    )
    doc.add_page_break()

    # TOC + Lists
    doc.add_paragraph("TABLE OF CONTENTS", style="Heading 1")
    _insert_toc(doc.add_paragraph())
    doc.add_page_break()

    doc.add_paragraph("LIST OF FIGURES", style="Heading 1")
    doc.add_paragraph("Figure 1.1 System Overview")
    doc.add_paragraph("Figure 3.1 Use Case Diagram")
    doc.add_paragraph("Figure 4.1 LangGraph Pipeline")
    doc.add_page_break()

    doc.add_paragraph("LIST OF TABLES", style="Heading 1")
    doc.add_paragraph("Table 3.1 Functional Requirements")
    doc.add_paragraph("Table 3.2 Non-Functional Requirements")
    doc.add_page_break()

def add_main_chapters(doc: Document):
    # New section for arabic numbering starting at 1
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(sec, 1.0)
    _set_page_number_format(sec, "decimal", 1)
    _add_page_number_footer(sec)

    # Chapter 1
    doc.add_paragraph("CHAPTER 1: INTRODUCTION", style="Heading 1")
    doc.add_paragraph("1.1 Overview", style="Heading 2")
    doc.add_paragraph(
        "The system enables natural-language logistics operations through an LLM planner inside a deterministic "
        "LangGraph. Planner selects intents; resolver maps entities; query/exec run tools; verify prompts for missing fields; "
        "reflect formats replies. WhatsApp via Twilio provides mobile-first operations."
    )
    doc.add_paragraph("1.2 Problem Statement", style="Heading 2")
    doc.add_paragraph(
        "Traditional systems are form-heavy, lose context, and reject partial inputs. Field staff prefer chat. "
        "We need conversational CRUD with context and deterministic execution."
    )
    doc.add_paragraph("1.3 Objectives", style="Heading 2")
    doc.add_paragraph(
        "- Conversational CRUD for drivers/vehicles/trips/loads/expenses/locations.\n"
        "- Context focus (e.g., trip expenses after trip details).\n"
        "- Deterministic tools; LLM only for planning.\n"
        "- WhatsApp integration with session memory."
    )
    doc.add_paragraph("1.4 Scope", style="Heading 2")
    doc.add_paragraph(
        "Includes agentic pipeline, tools, WhatsApp & web chat, testing harness. Excludes route optimization and payments."
    )
    doc.add_page_break()

    # Chapter 2
    doc.add_paragraph("CHAPTER 2: LITERATURE SURVEY", style="Heading 1")
    doc.add_paragraph("2.1 Existing Systems", style="Heading 2")
    doc.add_paragraph(
        "ERPs (SAP/Oracle) are comprehensive but costly; cloud SaaS is cheaper but form-centric; in-house systems are "
        "tailored but hard to maintain; AI-augmented tools are narrow. Natural language CRUD with context remains rare."
    )
    doc.add_paragraph("2.2 Comparative Analysis", style="Heading 2")
    doc.add_paragraph(
        "Our system combines conversational CRUD, context persistence, deterministic execution, and WhatsApp-native flow."
    )
    doc.add_paragraph("2.3 Research Gap", style="Heading 2")
    doc.add_paragraph(
        "Gaps: conversational operations, hybrid LLM+deterministic pipelines, multi-entity context, guided completion."
    )
    doc.add_page_break()

    # Chapter 3
    doc.add_paragraph("CHAPTER 3: SYSTEM ANALYSIS", style="Heading 1")
    doc.add_paragraph("3.1 Requirements", style="Heading 2")
    doc.add_paragraph(
        "Functional: register users/drivers, manage vehicles, trips, loads, expenses, locations; NL updates; entity "
        "resolution; focus memory. Non-functional: fast responses, high planner accuracy, deterministic tools."
    )
    doc.add_paragraph("3.2 Feasibility", style="Heading 2")
    doc.add_paragraph(
        "Technically/economically/operationally feasible with Python, LangGraph, SQLAlchemy, Twilio, and FastAPI."
    )
    doc.add_paragraph("3.3 Architecture Overview", style="Heading 2")
    doc.add_paragraph(
        "Nodes: router → planner → resolve → query/exec → verify → reflect, with optional loop when incomplete."
    )
    doc.add_page_break()

    # Chapter 4
    doc.add_paragraph("CHAPTER 4: SYSTEM DESIGN", style="Heading 1")
    doc.add_paragraph("4.1 LangGraph Pipeline", style="Heading 2")
    doc.add_paragraph(
        "Planner (LLM) selects intent/entities; resolver maps hints; exec/query run DB tools; verify produces questions; "
        "reflect formats; loop re-enters planner only when missing fields."
    )
    doc.add_paragraph("4.2 Database Design", style="Heading 2")
    doc.add_paragraph(
        "Entities: Owner, User/Driver, Vehicle, Trip, Load, Expense, LocationUpdate. Enums cover statuses."
    )
    doc.add_paragraph("4.3 Module Design", style="Heading 2")
    doc.add_paragraph(
        "Twilio FastAPI webhook for WhatsApp; web chat UI; test scenarios; NL update tool for single-field changes."
    )
    doc.add_page_break()

    # Chapter 5
    doc.add_paragraph("CHAPTER 5: IMPLEMENTATION", style="Heading 1")
    doc.add_paragraph("5.1 Technology Stack", style="Heading 2")
    doc.add_paragraph("Python, LangGraph, SQLAlchemy, FastAPI, Flask, Twilio, and LLM adapters.")
    doc.add_paragraph("5.2 Core Modules", style="Heading 2")
    doc.add_paragraph(
        "planner.py, resolve.py, query_agent.py, exec_mutation.py, verify.py, reflect.py, database_tools.py."
    )
    doc.add_paragraph("5.3 Integration", style="Heading 2")
    doc.add_paragraph("Twilio WhatsApp webhook + proactive sender; web chat with examples/help.")
    doc.add_page_break()

    # Chapter 6
    doc.add_paragraph("CHAPTER 6: TESTING AND VALIDATION", style="Heading 1")
    doc.add_paragraph("6.1 Strategy", style="Heading 2")
    doc.add_paragraph("Scenario harness covers queries, mutations, focus, NL updates, and missing fields.")
    doc.add_paragraph("6.2 Test Cases", style="Heading 2")
    doc.add_paragraph("12+ scenarios including trip expenses focus and load creation with missing fields.")
    doc.add_paragraph("6.3 Results", style="Heading 2")
    doc.add_paragraph("Most scenarios pass; NL update is robust; outputs are readable and structured.")
    doc.add_page_break()

    # Chapter 7
    doc.add_paragraph("CHAPTER 7: RESULTS AND DISCUSSION", style="Heading 1")
    doc.add_paragraph("7.1 Feature Outcomes", style="Heading 2")
    doc.add_paragraph("Conversational CRUD, focus-based follow-ups, clear missing-field prompts.")
    doc.add_paragraph("7.2 Performance", style="Heading 2")
    doc.add_paragraph("Fast-path queries are sub-second; planner used only when needed.")
    doc.add_paragraph("7.3 Limitations", style="Heading 2")
    doc.add_paragraph("No advanced route optimization; planner subject to model quotas.")
    doc.add_page_break()

    # Chapter 8
    doc.add_paragraph("CHAPTER 8: CONCLUSION AND FUTURE SCOPE", style="Heading 1")
    doc.add_paragraph("8.1 Conclusion", style="Heading 2")
    doc.add_paragraph(
        "Hybrid LLM + deterministic orchestration enables reliable, usable logistics operations via chat/WhatsApp."
    )
    doc.add_paragraph("8.2 Future Enhancements", style="Heading 2")
    doc.add_paragraph("Route optimization, analytics, multi-tenant RBAC, voice interface, richer WhatsApp templates.")
    doc.add_page_break()

    # References
    doc.add_paragraph("REFERENCES", style="Heading 1")
    for ref in [
        "LangGraph Docs: https://langchain-ai.github.io/langgraph/",
        "Twilio WhatsApp: https://www.twilio.com/whatsapp",
        "SQLAlchemy: https://www.sqlalchemy.org/",
        "FastAPI: https://fastapi.tiangolo.com/",
    ]:
        doc.add_paragraph(ref)
    doc.add_page_break()

    # Appendix
    doc.add_paragraph("APPENDIX", style="Heading 1")
    doc.add_paragraph("Sample commands and API endpoints omitted for brevity.")

def main():
    doc = Document()
    set_base_styles(doc)
    add_cover(doc)
    add_certificate_decl_frontmatter(doc)
    add_main_chapters(doc)
    doc.save("PBL-IV REPORT.docx")

if __name__ == "__main__":
    main()